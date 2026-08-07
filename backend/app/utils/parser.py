"""简历文件解析器 — 支持 PDF/DOCX/Markdown 格式"""

import re
from pathlib import Path
from typing import Optional


class ResumeParseResult:
    """解析结果"""
    def __init__(self, raw_text: str, source_format: str, source_filename: str,
                 paragraphs: Optional[list[dict]] = None):
        self.raw_text = raw_text
        self.source_format = source_format
        self.source_filename = source_filename
        self.paragraphs = paragraphs or []  # [{text, x0, y0, x1, y1, font_size, is_bold}]

    def __repr__(self):
        return f"<ResumeParseResult(format={self.source_format}, chars={len(self.raw_text)})>"


def parse_resume(file_path: str) -> ResumeParseResult:
    """根据文件扩展名自动选择解析器"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(file_path)
    elif suffix == ".docx":
        return parse_docx(file_path)
    elif suffix in (".md", ".markdown"):
        return parse_markdown(file_path)
    elif suffix == ".txt":
        return parse_text(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}。支持: PDF, DOCX, Markdown, TXT")


def parse_pdf(file_path: str) -> ResumeParseResult:
    """解析 PDF 简历，提取文本和坐标信息"""
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    paragraphs = []
    full_text_parts = []

    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block["type"] != 0:  # 跳过图片等非文本块
                continue

            for line in block["lines"]:
                line_text = ""
                line_bbox = None
                font_sizes = []
                is_bold = False

                for span in line["spans"]:
                    line_text += span["text"]
                    font_sizes.append(span["size"])
                    if "Bold" in span.get("font", ""):
                        is_bold = True
                    # 取第一个span的边界框
                    if line_bbox is None:
                        line_bbox = list(span["bbox"])

                line_text = line_text.strip()
                if not line_text:
                    continue

                paragraphs.append({
                    "text": line_text,
                    "x0": line_bbox[0] if line_bbox else 0,
                    "y0": line_bbox[1] if line_bbox else 0,
                    "x1": line_bbox[2] if line_bbox else 0,
                    "y1": line_bbox[3] if line_bbox else 0,
                    "font_size": max(font_sizes) if font_sizes else 0,
                    "is_bold": is_bold,
                    "page": page_num,
                })
                full_text_parts.append(line_text)

    doc.close()
    raw_text = "\n".join(full_text_parts)

    filename = Path(file_path).name
    return ResumeParseResult(
        raw_text=raw_text,
        source_format="pdf",
        source_filename=filename,
        paragraphs=paragraphs,
    )


def parse_docx(file_path: str) -> ResumeParseResult:
    """解析 DOCX 简历"""
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    full_text_parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # 保留空行以维持段落间距
            full_text_parts.append("")
            continue

        # 检测粗体 (可能用于标题)
        is_bold = any(run.bold for run in para.runs if run.bold)
        # 检测字体大小
        font_sizes = [run.font.size for run in para.runs if run.font.size]
        font_size = max(font_sizes).pt if font_sizes else 0

        paragraphs.append({
            "text": text,
            "font_size": font_size,
            "is_bold": is_bold,
            "style": para.style.name if para.style else "",
        })
        full_text_parts.append(text)

    raw_text = "\n".join(full_text_parts)

    filename = Path(file_path).name
    return ResumeParseResult(
        raw_text=raw_text,
        source_format="docx",
        source_filename=filename,
        paragraphs=paragraphs,
    )


def parse_markdown(file_path: str) -> ResumeParseResult:
    """解析 Markdown 简历"""
    with open(file_path, "r", encoding="utf-8") as f:
        raw_text = f.read()

    # 简单按标题分割段落
    lines = raw_text.split("\n")
    paragraphs = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        is_heading = stripped.startswith("#")
        font_size = 20 if is_heading else 12
        paragraphs.append({
            "text": stripped,
            "font_size": font_size,
            "is_bold": is_heading,
            "is_heading": is_heading,
        })

    filename = Path(file_path).name
    return ResumeParseResult(
        raw_text=raw_text,
        source_format="md",
        source_filename=filename,
        paragraphs=paragraphs,
    )


def parse_text(file_path: str) -> ResumeParseResult:
    """解析纯文本简历"""
    with open(file_path, "r", encoding="utf-8") as f:
        raw_text = f.read()

    filename = Path(file_path).name
    return ResumeParseResult(
        raw_text=raw_text,
        source_format="txt",
        source_filename=filename,
    )


def reconstruct_reading_order(paragraphs: list[dict]) -> list[dict]:
    """基于坐标的阅读顺序重建 (PDF专用)

    解决双栏布局、表格等复杂排版导致的文本提取顺序错乱问题。
    """
    if not paragraphs or "y0" not in paragraphs[0]:
        return paragraphs

    # 按 y 坐标排序 (从上到下)，同一行内按 x 坐标排序 (从左到右)
    # 同一行的判定: y0 坐标差距在 10 个像素以内
    sorted_paragraphs = sorted(paragraphs, key=lambda p: (p.get("page", 0), p.get("y0", 0), p.get("x0", 0)))

    result = []
    current_line = []
    current_y = None

    for p in sorted_paragraphs:
        y0 = p.get("y0", 0)
        if current_y is None or abs(y0 - current_y) <= 10:
            current_line.append(p)
            current_y = y0 if current_y is None else current_y
        else:
            # 同一行内按 x 排序
            current_line.sort(key=lambda x: x.get("x0", 0))
            result.extend(current_line)
            current_line = [p]
            current_y = y0

    # 处理最后一行
    if current_line:
        current_line.sort(key=lambda x: x.get("x0", 0))
        result.extend(current_line)

    return result
